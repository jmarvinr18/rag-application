#!/usr/bin/env python3
"""
confluence_pdf_to_docx.py
=========================

Convert a Confluence-exported PDF into a single, self-contained .docx and
apply the requested cleanup. Images are embedded *inside* the .docx (not
referenced externally), so the one file carries everything.

Built to run unattended in Jenkins: pure CLI, no prompts, clear exit codes.

Cleanup steps:
  1. Identify headers/subheaders -> Word heading styles (Heading 1/2/3...),
                                    which also drive Word's navigation outline.
  2. Trim whitespace             -> collapses repeated spaces, NBSP and
                                    zero-width chars, drops empty paragraphs.
  3. Remove divider lines        -> rows of dashes/underscores are skipped;
                                    drawn rules carry no text, so they vanish.
  4. Hyperlinks show exact URL   -> the anchor is rendered as the real URL,
                                    as a true clickable Word hyperlink.
  5. Preserve images             -> embedded inline in the .docx via add_picture.

Exit codes (for Jenkins):
  0  success
  1  failure

Usage:
  pip install pymupdf python-docx
  python confluence_pdf_to_docx.py input.pdf output.docx
  python confluence_pdf_to_docx.py input.pdf output.docx --min-image-size 24 --quiet
"""

import argparse
import io
import re
import sys
from collections import Counter

try:
    import fitz  # PyMuPDF
except ImportError:
    sys.exit("PyMuPDF is required. Install it with:  pip install pymupdf")

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
except ImportError:
    sys.exit("python-docx is required. Install it with:  pip install python-docx")


# --- tunables -------------------------------------------------------------- #
HEADER_SIZE_RATIO = 1.15
MAX_HEADER_LEVELS = 4
MAX_HEADER_CHARS = 140
MAX_IMAGE_WIDTH_IN = 6.5            # cap image width to the content area
WORD_SAFE_IMG = {"png", "jpeg", "jpg", "bmp", "gif", "tiff", "tif"}
DIVIDER_RE = re.compile(r"^\s*[-_=*~–—•·.]{3,}\s*$")
_WS_RE = re.compile(r"[ \t\f\v]+")
_ZERO_WIDTH = ("\u200b", "\u200c", "\u200d", "\ufeff")


def clean_ws(text: str) -> str:
    text = text.replace("\xa0", " ")
    for zw in _ZERO_WIDTH:
        text = text.replace(zw, "")
    return _WS_RE.sub(" ", text).strip()


# --- font-size analysis ---------------------------------------------------- #
def compute_body_size(doc):
    sizes = Counter()
    for page in doc:
        for block in page.get_text("dict")["blocks"]:
            if block.get("type") != 0:
                continue
            for line in block["lines"]:
                for span in line["spans"]:
                    if span["text"].strip():
                        sizes[round(span["size"], 1)] += len(span["text"].strip())
    return sizes.most_common(1)[0][0] if sizes else 0.0


def build_header_levels(doc, body_size):
    big = set()
    for page in doc:
        for block in page.get_text("dict")["blocks"]:
            if block.get("type") != 0:
                continue
            for line in block["lines"]:
                for span in line["spans"]:
                    size = round(span["size"], 1)
                    if size > body_size * HEADER_SIZE_RATIO and span["text"].strip():
                        big.add(size)
    ranked = sorted(big, reverse=True)
    return {s: min(i + 1, MAX_HEADER_LEVELS) for i, s in enumerate(ranked)}


def dominant_size(line):
    sizes = Counter()
    for span in line["spans"]:
        sizes[round(span["size"], 1)] += max(1, len(span["text"].strip()))
    return sizes.most_common(1)[0][0] if sizes else 0.0


# --- hyperlinks ------------------------------------------------------------ #
def uri_for_span(span_bbox, link_rects):
    srect = fitz.Rect(span_bbox)
    for rect, uri in link_rects:
        if srect.intersects(rect):
            return uri
    return None


def page_link_rects(page):
    out = []
    for link in page.get_links():
        uri = link.get("uri")
        if uri and "from" in link:
            out.append((fitz.Rect(link["from"]), uri))
    return out


def line_segments(line, link_rects):
    """Split a line into ('text', str) and ('link', uri) segments."""
    segs = []
    prev_uri = None
    for span in line["spans"]:
        uri = uri_for_span(span["bbox"], link_rects)
        if uri:
            if uri != prev_uri:
                segs.append(("link", uri))
            prev_uri = uri
        else:
            if segs and segs[-1][0] == "text":
                segs[-1] = ("text", segs[-1][1] + span["text"])
            else:
                segs.append(("text", span["text"]))
            prev_uri = None
    return segs


def add_hyperlink(paragraph, url, text):
    """Insert a real, clickable Word hyperlink whose display text is the URL."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    color = OxmlElement("w:color"); color.set(qn("w:val"), "0563C1"); rpr.append(color)
    underline = OxmlElement("w:u"); underline.set(qn("w:val"), "single"); rpr.append(underline)
    run.append(rpr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


# --- images ---------------------------------------------------------------- #
def image_bytes_for_word(block):
    """Return (png_or_orig_bytes, ext) that Word can embed; convert if needed."""
    data = block.get("image")
    if not data:
        return None, None
    ext = (block.get("ext") or "png").lower()
    if ext in WORD_SAFE_IMG:
        return data, ext
    try:  # exotic format (jpx, jbig2, cmyk...) -> re-encode to PNG
        pix = fitz.Pixmap(data)
        if pix.alpha or (pix.colorspace and pix.colorspace.n > 3):
            pix = fitz.Pixmap(fitz.csRGB, pix)
        return pix.tobytes("png"), "png"
    except Exception:
        return None, None


# --- page -> ordered elements --------------------------------------------- #
def page_elements(page, level_map, min_size):
    """Yield ('heading', lvl, text) | ('para', segs) | ('image', bytes, w_in)."""
    link_rects = page_link_rects(page)
    elements = []
    for block in page.get_text("dict")["blocks"]:
        x0, y0 = block["bbox"][0], block["bbox"][1]

        if block.get("type") == 1:  # image
            bbox = fitz.Rect(block["bbox"])
            if min_size and (bbox.width < min_size or bbox.height < min_size):
                continue
            data, _ = image_bytes_for_word(block)
            if data:
                width_in = min(bbox.width / 72.0, MAX_IMAGE_WIDTH_IN)
                elements.append((y0, x0, ("image", data, max(width_in, 0.3))))
            continue

        if block.get("type") != 0:
            continue

        para_segs = []
        for line in block["lines"]:
            raw = "".join(s["text"] for s in line["spans"])
            if DIVIDER_RE.match(clean_ws(raw)):     # step 3: drop dividers
                continue
            text = clean_ws(raw)
            if not text:
                continue
            size = dominant_size(line)
            level = level_map.get(size)
            if level and len(text) <= MAX_HEADER_CHARS:    # step 1: heading
                if para_segs:
                    elements.append((y0, x0, ("para", para_segs)))
                    para_segs = []
                ly, lx = line["bbox"][1], line["bbox"][0]
                elements.append((ly, lx, ("heading", level, text)))
            else:
                if para_segs:
                    para_segs.append(("text", " "))
                para_segs.extend(line_segments(line, link_rects))

        if para_segs:
            elements.append((y0, x0, ("para", para_segs)))

    elements.sort(key=lambda e: (round(e[0], 1), round(e[1], 1)))
    return [el for _, _, el in elements]


# --- docx builder ---------------------------------------------------------- #
def render_paragraph(document, segs):
    # normalise whitespace and merge adjacent text segments
    norm = []
    for kind, val in segs:
        if kind == "text":
            val = clean_ws(val)
            if not val:
                continue
            if norm and norm[-1][0] == "text":
                norm[-1] = ("text", norm[-1][1] + " " + val)
            else:
                norm.append(("text", val))
        else:
            norm.append(("link", val))
    if not norm:
        return
    p = document.add_paragraph()
    for i, (kind, val) in enumerate(norm):
        if i > 0:
            p.add_run(" ")
        if kind == "text":
            p.add_run(val)
        else:
            add_hyperlink(p, val, val)   # step 4: display the exact URL


def convert(input_path, output_path, min_size, quiet):
    src = fitz.open(input_path)
    if src.page_count == 0:
        raise ValueError("PDF has no pages")

    body_size = compute_body_size(src)
    level_map = build_header_levels(src, body_size)

    document = Document()
    # python-docx ships a default settings.xml whose <w:zoom> lacks the
    # required w:percent attribute; drop it so the file validates strictly.
    settings = document.settings.element
    for zoom in settings.findall(qn("w:zoom")):
        settings.remove(zoom)

    counts = {"headings": 0, "paras": 0, "links": 0, "images": 0}

    for page in src:
        for el in page_elements(page, level_map, min_size):
            if el[0] == "heading":
                document.add_heading(el[2], level=el[1])
                counts["headings"] += 1
            elif el[0] == "para":
                before = counts["links"]
                counts["links"] += sum(1 for k, _ in el[1] if k == "link")
                render_paragraph(document, el[1])
                counts["paras"] += 1
                _ = before
            elif el[0] == "image":
                try:
                    document.add_picture(io.BytesIO(el[1]), width=Inches(el[2]))
                    counts["images"] += 1
                except Exception as exc:
                    if not quiet:
                        print(f"  (skipped one image: {exc})", file=sys.stderr)

    document.save(output_path)
    src.close()

    if not quiet:
        print("Conversion complete.")
        print(f"  Headings   : {counts['headings']}")
        print(f"  Paragraphs : {counts['paras']}")
        print(f"  Hyperlinks : {counts['links']}")
        print(f"  Images     : {counts['images']}")
        print(f"  Output     : {output_path}")


def main():
    ap = argparse.ArgumentParser(
        description="Convert a Confluence PDF to a self-contained .docx for Amazon Q."
    )
    ap.add_argument("input", help="Source PDF")
    ap.add_argument("output", help="Destination .docx")
    ap.add_argument("--min-image-size", type=float, default=0,
                    help="Skip images smaller than this many points (drops logos/icons).")
    ap.add_argument("--quiet", action="store_true", help="Suppress the summary")
    args = ap.parse_args()

    try:
        convert(args.input, args.output, args.min_image_size, args.quiet)
    except Exception as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        sys.exit(1)
    sys.exit(0)


if __name__ == "__main__":
    main()
